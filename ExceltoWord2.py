#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import argparse
import os
import re
import sys
from typing import List, Tuple, Optional

import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING


# ---------- constants ----------
EN_DASH = "–"          # U+2013
GREEK_ALPHA = "α"      # U+03B1
UNICODE_MINUS = "−"    # U+2212


# ---------- text pre-processing (length-preserving) ----------
def preprocess_fragment(s: str) -> str:
    if not s:
        return ""
    # digit-hyphen-digit -> en dash
    s = re.sub(r'(?<=\d)-(?=\d)', EN_DASH, s)
    # "δ , " -> "δ "
    s = s.replace("δ , ", "δ ")
    return s


# ---------- styling engine: produce runs with flags ----------
def build_styled_runs(text: str, title_len: int):
    """
    Returns list of runs: (substr, bold, italic, superscript, subscript).
    Title is the first title_len characters.
    """
    n = len(text)
    if n == 0:
        return []

    bold = [False] * n
    italic = [False] * n
    sup = [False] * n
    sub = [False] * n

    # bold for title
    for i in range(min(title_len, n)):
        bold[i] = True

    # Title-only: [α]D -> D subscript; digits immediately after (until space) superscript
    if title_len > 0:
        i = 0
        while i < min(title_len, n):
            if text.startswith(f"[{GREEK_ALPHA}]D", i):
                d_idx = i + len(f"[{GREEK_ALPHA}]D") - 1
                if d_idx < n:
                    sub[d_idx] = True
                j = d_idx + 1
                while j < min(title_len, n) and text[j].isdigit():
                    sup[j] = True
                    j += 1
                i = j
                continue
            i += 1

    # Title-only: nucleus numbers before H/C/F superscript
    if title_len > 0:
        i = 0
        while i < min(title_len, n):
            ch = text[i]
            if ch in ("H", "C", "F") and i > 0:
                j = i - 1
                touched = False
                while j >= 0 and j < title_len and text[j].isdigit():
                    sup[j] = True
                    touched = True
                    j -= 1
                if touched:
                    i += 1
                    continue
            i += 1

    # Line-wide: italicize "m/z"
    for m in re.finditer(r"m/z", text):
        for k in range(m.start(), m.end()):
            italic[k] = True

    # Line-wide: italicize J in "J =" or "J="
    for m in re.finditer(r"J\s*=", text):
        italic[m.start()] = True  # only the 'J'

    # Line-wide: [α]D in body -> D subscript (harmless if already done in title)
    i = 0
    while i < n:
        if text.startswith(f"[{GREEK_ALPHA}]D", i):
            d_idx = i + len(f"[{GREEK_ALPHA}]D") - 1
            if d_idx < n:
                sub[d_idx] = True
            i = d_idx + 1
            continue
        i += 1

    # Line-wide: ']' followed by digits/± -> superscript (e.g., [M+Na]+, [M]2+, [M−H]−)
    i = 0
    while i < n:
        if text[i] == "]":
            j = i + 1
            while j < n and (text[j].isdigit() or text[j] in "+-" or text[j] == UNICODE_MINUS):
                sup[j] = True
                j += 1
            i = j
            continue
        i += 1

    # Line-wide: subscript full digit runs that immediately follow ASCII letters
    # (formulas/solvents/-d6/-d8). Skip if digits already superscripted.
    i = 0
    while i < n:
        if text[i].isalpha() and i + 1 < n and text[i + 1].isdigit():
            j = i + 1
            while j < n and text[j].isdigit():
                j += 1
            if not any(sup[k] for k in range(i + 1, j)):
                for k in range(i + 1, j):
                    sub[k] = True
            i = j
            continue
        i += 1

    # Collapse masks to runs (superscript beats subscript if both True)
    def key(idx: int):
        sb = bold[idx]
        it = italic[idx]
        su = sup[idx]
        sb2 = sub[idx] and not su
        return (sb, it, su, sb2)

    runs = []
    cur = key(0)
    buf = [text[0]]
    for i in range(1, n):
        if key(i) == cur:
            buf.append(text[i])
        else:
            runs.append(("".join(buf), *cur))
            buf = [text[i]]
            cur = key(i)
    runs.append(("".join(buf), *cur))
    return runs


# ---------- paragraph writer ----------
def write_paragraph(doc: Document, title: str, body: str) -> None:
    """Create one SI line with all styles applied.
       NOTE: no space before colon if body starts with ':' or '：'."""
    title2 = preprocess_fragment(title or "")
    body2 = preprocess_fragment(body or "")

    # smart joiner (no space before colon)
    if body2:
        joiner = "" if body2.startswith(":") or body2.startswith("：") else " "
        whole = f"{title2}{joiner}{body2}"
    else:
        whole = title2

    title_len = len(title2)
    runs = build_styled_runs(whole, title_len)

    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(16)

    for text, is_bold, is_italic, is_sup, is_sub in runs:
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
        r.bold = is_bold
        r.italic = is_italic
        r.font.superscript = is_sup
        r.font.subscript = is_sub


# ---------- Excel ingestion ----------
def read_pairs_from_excel(xlsx_path: str, sheet: str = "Python") -> List[Tuple[str, str]]:
    """Read (title, body) pairs from sheet 'Python' supporting 2×N or N×2."""
    try:
        df = pd.read_excel(xlsx_path, sheet_name=sheet, header=None, dtype=str)
    except Exception as e:
        raise RuntimeError(f"Failed reading sheet '{sheet}': {e}")

    df = df.dropna(how="all").dropna(axis=1, how="all")
    if df.empty:
        raise RuntimeError("Sheet is empty after removing blank rows/columns.")
    df = df.fillna("")

    rows, cols = df.shape
    pairs: List[Tuple[str, str]] = []

    if rows == 2 and cols >= 1:
        for c in range(cols):
            t = str(df.iat[0, c]).strip()
            b = str(df.iat[1, c]).strip()
            if t or b:
                pairs.append((t, b))
    elif cols == 2 and rows >= 1:
        for r in range(rows):
            t = str(df.iat[r, 0]).strip()
            b = str(df.iat[r, 1]).strip()
            if t or b:
                pairs.append((t, b))
    else:
        sub = df.iloc[:2, :5] if rows >= 2 and cols >= 1 else df
        sub = sub.fillna("")
        if sub.shape[0] == 2:
            for c in range(sub.shape[1]):
                t = str(sub.iat[0, c]).strip()
                b = str(sub.iat[1, c]).strip()
                if t or b:
                    pairs.append((t, b))
        elif sub.shape[1] == 2:
            for r in range(sub.shape[0]):
                t = str(sub.iat[r, 0]).strip()
                b = str(sub.iat[r, 1]).strip()
                if t or b:
                    pairs.append((t, b))
        else:
            raise RuntimeError(f"Unrecognized layout (rows={rows}, cols={cols}); expected 2×N or N×2.")

    return pairs


# ---------- NMR/HRMS parsing for consistency check ----------
def parse_h1_total_from_body(h1_body: str) -> Optional[int]:
    """Sum of all 'nH' occurrences in ¹H NMR body, ignoring 'Hz'."""
    if not h1_body:
        return None
    nums = [int(x) for x in re.findall(r'(\d+)\s*H(?!z)', h1_body, flags=re.IGNORECASE)]
    return sum(nums) if nums else 0


def parse_c13_count_from_body(c13_body: str) -> Optional[int]:
    """
    Sum ¹³C NMR carbon counts.

    Each comma-separated 13C shift is counted as 1 carbon unless it has an
    explicit multiplier such as (2C), (4C), or (14C), in which case that
    multiplier is added. This mirrors the ¹H parser behavior, which sums nH
    rather than merely counting comma-separated signals.
    """
    if not c13_body:
        return None

    s = c13_body.replace("δ", " ").replace("−", "-").replace("–", "-")

    # Split on commas only when the comma is not inside parentheses. This keeps
    # annotations such as "(d, J = ... Hz)" in a single segment.
    segments = re.split(r',\s*(?![^()]*\))', s)

    total = 0
    found_shift = False

    for seg in segments:
        seg = seg.strip()
        if not seg:
            continue

        # A 13C chemical shift, optionally written as a range.
        # Examples: 130.1, 129.92-129.77, 129.92–129.77
        if not re.search(r'(?<![A-Za-z0-9.])\d{1,3}(?:\.\d+)?(?:\s*-\s*\d{1,3}(?:\.\d+)?)?', seg):
            continue

        found_shift = True

        # Explicit carbon-count multiplier, e.g., (2C), (14 C).
        m = re.search(r'\((\d+)\s*C\)', seg, flags=re.IGNORECASE)
        if m:
            total += int(m.group(1))
        else:
            total += 1

    return total if found_shift else 0


def parse_hrms_CH_from_line(hrms_text: str) -> Tuple[Optional[int], Optional[int]]:
    """Extract C and H counts from a molecular formula in an HRMS line."""
    if not hrms_text:
        return (None, None)
    # Try compact Hill-style formula first: C43H65..., then forgiving spaced variants.
    m = re.search(r'C(\d+)H(\d+)', hrms_text)
    if not m:
        m = re.search(r'C\s*(\d+)\s*H\s*(\d+)', hrms_text)
    if m:
        return (int(m.group(1)), int(m.group(2)))
    return (None, None)


def extract_relevant_texts(pairs: List[Tuple[str, str]]) -> Tuple[Optional[str], Optional[str], Optional[str]]:
    """Return the bodies (or combined line) for 1H, 13C, and HRMS respectively."""
    h1_body = None
    c13_body = None
    hrms_line = None
    for title, body in pairs:
        t = (title or "").strip()
        tU = t.upper()
        if h1_body is None and re.match(r'^\s*1H\b', tU):
            h1_body = body or ""
        elif c13_body is None and re.match(r'^\s*13C', tU):
            c13_body = body or ""
        elif hrms_line is None and tU.startswith("HRMS"):
            # combine title+body to maximize chance to see the formula
            joiner = "" if (body or "").startswith(":") or (body or "").startswith("：") else " "
            hrms_line = f"{title}{joiner}{body}"
    return h1_body, c13_body, hrms_line


def append_consistency_check(doc: Document, pairs: List[Tuple[str, str]]) -> None:
    """Compute ¹H/¹³C from NMR and compare with HRMS C/H; append a short report."""
    h1_body, c13_body, hrms_line = extract_relevant_texts(pairs)

    h1_sum = parse_h1_total_from_body(h1_body) if h1_body is not None else None
    c13_cnt = parse_c13_count_from_body(c13_body) if c13_body is not None else None
    c_ms, h_ms = parse_hrms_CH_from_line(hrms_line or "")

    # Build a compact report
    lines = []
    lines.append("Consistency check — NMR vs HRMS")

    def verdict(nmr_val, ms_val, label):
        if nmr_val is None or ms_val is None:
            return f"{label}: n/a"
        delta = ms_val - nmr_val
        if delta == 0:
            return f"{label}: OK (Δ = 0)"
        sign = "+" if delta > 0 else "−"
        return f"{label}: mismatch (Δ = {sign}{abs(delta)})"

    # Use Unicode superscripts for ¹H / ¹³C labels
    h_line = f"¹H total (from 1H NMR) = {h1_sum if h1_sum is not None else 'n/a'}; HRMS H = {h_ms if h_ms is not None else 'n/a'} — {verdict(h1_sum, h_ms, 'H')}"
    c_line = f"¹³C count (from 13C NMR) = {c13_cnt if c13_cnt is not None else 'n/a'}; HRMS C = {c_ms if c_ms is not None else 'n/a'} — {verdict(c13_cnt, c_ms, 'C')}"

    lines.append(h_line)
    lines.append(c_line)

    # Append to document (one paragraph per line)
    for text in [""] + lines:  # leading blank line
        p = doc.add_paragraph(text)
        pf = p.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(16)
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)


# ---------- main ----------
def main():
    ap = argparse.ArgumentParser(
        description="Read Excel sheet 'Python' (A1:E2; 2×N or N×2) and produce a formatted .docx SI section."
    )
    ap.add_argument("excel", nargs="?", help="Path to Excel workbook")
    ap.add_argument("-s", "--sheet", default="Python", help="Worksheet name (default: Python)")
    ap.add_argument("-o", "--output", help="Output .docx path (default: <excel_basename>_SI.docx)")
    args = ap.parse_args()

    xlsx = args.excel or input("Excel file path: ").strip()
    if not os.path.isfile(xlsx):
        print(f"Error: file not found: {xlsx}", file=sys.stderr)
        sys.exit(1)

    try:
        pairs = read_pairs_from_excel(xlsx, sheet=args.sheet)
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(2)

    doc = Document()
    # Base font on Normal style
    try:
        style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)
    except Exception:
        pass

    for title, body in pairs:
        write_paragraph(doc, title, body)

    # NEW: append ¹H/¹³C vs HRMS consistency check
    append_consistency_check(doc, pairs)

    out_path = args.output if args.output else os.path.splitext(xlsx)[0] + "_SI.docx"
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    main()
